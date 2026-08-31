import json
import re
from collections import defaultdict
from datetime import date, datetime
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document as WordDocument
from fastapi import HTTPException, status
from openpyxl import load_workbook
from openpyxl.utils import column_index_from_string, get_column_letter
from sqlalchemy import select
from sqlalchemy.orm import Session

from app.models import CandidateStatus, FieldProvenance, ImportBatch, ImportMatchReview, ImportStatus, RelatedInfoCandidate, SourceDocument, Student, StudentRelatedInfoCard, User, WordImportCandidate, utcnow
from app.schemas import CandidateApproval
from app.services.ai import extract_related_info, extract_word_candidates
from app.services.crypto import blind_index
from app.services.students import STUDENT_FIELDS, record_student_version, student_to_dict

HEADER_ALIASES = {
    "student_no": {"学号", "学生编号", "学生id", "studentid", "studentno", "studentnumber", "id"},
    "candidate_no": {"考生号", "考号", "准考证号", "candidateno", "examno"},
    "full_name": {"姓名", "中文姓名", "学生姓名", "name", "fullname"},
    "gender": {"性别", "gender"},
    "national_id": {"身份证号", "身份证号码", "身份证", "idcard", "nationalid"},
    "date_of_birth": {"出生日期", "生日", "dateofbirth", "birthdate"},
    "student_origin": {"生源地", "生源所在地", "studentorigin"},
    "ethnicity": {"民族", "ethnicity", "nation"},
    "political_status": {"政治面貌", "politicalstatus"},
    "enrollment_date": {"入学日期", "入学时间", "enrollmentdate"},
    "graduation_year": {"毕业年份", "毕业年", "graduationyear"},
    "graduation_date": {"毕业日期", "毕业时间", "graduationdate"},
    "urban_rural_origin": {"城乡生源", "城乡类型", "urbanruralorigin"},
    "pre_enrollment_archive_unit": {"入学前档案所在单位", "档案所在单位", "preenrollmentarchiveunit"},
    "archive_transferred": {"档案是否转入学校", "档案转入学校", "archivetransferred"},
    "pre_enrollment_police_station": {"入学前户口所在地派出所", "户口所在地派出所", "preenrollmentpolicestation"},
    "household_registration_transferred": {"户口是否转入学校", "户口转入学校", "householdregistrationtransferred"},
    "education_level": {"学历层次", "学历", "educationlevel"},
    "program_duration": {"学制", "学制年限", "programduration"},
    "school": {"所属学校", "学校", "school"},
    "college": {"所属学院", "学院", "college", "faculty"},
    "school_major": {"学校专业", "专业", "所学专业", "major", "specialty"},
    "major_direction": {"专业方向", "方向", "majordirection"},
    "current_class": {"所在班级", "班级", "class", "classname"},
    "training_mode": {"培养方式", "trainingmode"},
    "commissioned_unit": {"委培单位", "委托培养单位", "commissionedunit"},
    "hardship_category": {"困难生类别", "困难类别", "hardshipcategory"},
    "normal_student_category": {"师范生类别", "师范类别", "normalstudentcategory"},
    "mobile_phone": {"手机号码", "手机", "电话", "手机号", "联系电话", "phone", "mobile"},
    "electronic_email": {"邮箱", "电子邮箱", "email"},
    "qq_number": {"qq号码", "qq号", "qq", "qqnumber"},
    "family_phone": {"家庭电话", "家庭联系电话", "familyphone"},
    "family_postcode": {"家庭邮编", "家庭邮政编码", "familypostcode"},
    "family_address": {"家庭地址", "地址", "家庭住址", "address"},
    "poverty_county_52": {"是否52个贫困县", "是否52个贫困县生源", "povertycounty52"},
    "poverty_county_province": {"贫困县所在省", "贫困县省", "povertycountyprovince"},
    "poverty_county_city": {"贫困县所在市", "贫困县市", "povertycountycity"},
    "poverty_county_district": {"贫困县所在县", "贫困县县", "povertycountydistrict"},
    "registered_poor": {"是否建档立卡", "建档立卡", "registeredpoor"},
    "study_mode": {"学习形式", "学习方式", "studymode"},
    "vocational_expansion_flag": {"高职扩招考生标志", "高职扩招标志", "vocationalexpansionflag"},
    "remarks": {"备注", "说明", "备注信息", "remarks", "notes"},
}

DATE_FIELDS = {"date_of_birth", "enrollment_date", "graduation_date"}
IMPORT_MODES = {"upsert", "create_only", "update_only"}
UPDATE_POLICIES = {"overwrite", "only_blank"}


def _normalize_header(value: Any) -> str:
    return re.sub(r"[\s_\-()（）]", "", str(value or "").strip().lower())


def _as_text(value: Any) -> str | None:
    if value is None:
        return None
    if isinstance(value, float) and value.is_integer():
        return str(int(value))
    if isinstance(value, datetime):
        return value.date().isoformat()
    if isinstance(value, date):
        return value.isoformat()
    value = str(value).strip()
    return value or None


def _as_date(value: Any) -> date | None:
    if not value:
        return None
    if isinstance(value, datetime):
        return value.date()
    if isinstance(value, date):
        return value
    text = _as_text(value)
    if not text:
        return None
    for pattern in ("%Y-%m-%d", "%Y/%m/%d", "%Y.%m.%d", "%Y%m%d", "%Y-%m", "%Y/%m", "%Y.%m", "%Y%m"):
        try:
            return datetime.strptime(text, pattern).date()
        except ValueError:
            continue
    raise ValueError("日期格式应为 YYYY-MM-DD")


def _map_headers(headers: list[Any]) -> dict[int, str]:
    mapping: dict[int, str] = {}
    normalized_aliases = {field: {_normalize_header(item) for item in aliases} for field, aliases in HEADER_ALIASES.items()}
    for column_index, header in enumerate(headers, start=1):
        normalized = _normalize_header(header)
        for field, aliases in normalized_aliases.items():
            if normalized in aliases:
                mapping[column_index] = field
                break
    return mapping


def _open_student_worksheet(content: bytes):
    try:
        workbook = load_workbook(BytesIO(content), read_only=True, data_only=True)
    except Exception as exc:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="无法读取 Excel 文件") from exc
    worksheet = next((sheet for sheet in workbook.worksheets if sheet.sheet_state == "visible"), None)
    if worksheet is None:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Excel 没有可读取的工作表")
    headers = next(worksheet.iter_rows(min_row=1, max_row=1, values_only=True), None)
    if not headers:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Excel 缺少表头")
    return workbook, worksheet, list(headers)


def _normalize_mapping(headers: list[Any], requested_mapping: dict[str, Any] | None = None) -> dict[int, str]:
    if not requested_mapping:
        mapping = _map_headers(headers)
    else:
        mapping = {}
        for raw_column, raw_field in requested_mapping.items():
            field = str(raw_field or "").strip()
            if not field or field not in STUDENT_FIELDS:
                continue
            try:
                if isinstance(raw_column, int) or str(raw_column).isdigit():
                    column = int(raw_column)
                else:
                    column = column_index_from_string(str(raw_column).strip().upper())
            except ValueError:
                continue
            if 1 <= column <= len(headers) and field not in mapping.values():
                mapping[column] = field
    if "student_no" not in mapping.values() or "full_name" not in mapping.values():
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="字段映射至少需要“学号”和“姓名”列")
    return mapping


def _parsed_excel_fields(values: tuple[Any, ...], mapping: dict[int, str]) -> tuple[dict[str, Any], dict[str, Any]]:
    raw_fields = {field: values[column - 1] if column <= len(values) else None for column, field in mapping.items()}
    parsed_fields: dict[str, Any] = {}
    for field, raw_value in raw_fields.items():
        value = _as_date(raw_value) if field in DATE_FIELDS else _as_text(raw_value)
        if value is not None:
            parsed_fields[field] = value
    return raw_fields, parsed_fields


def _display_import_value(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, (date, datetime)):
        return value.isoformat()
    return str(value).strip()


def preview_excel_import(
    db: Session,
    document: SourceDocument,
    content: bytes,
    mode: str = "upsert",
    requested_mapping: dict[str, Any] | None = None,
    required_fields: list[str] | None = None,
    update_policy: str = "overwrite",
) -> tuple[dict[str, Any], dict[int, str]]:
    if mode not in IMPORT_MODES:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="无效的导入模式")
    if update_policy not in UPDATE_POLICIES:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="无效的更新策略")
    workbook, worksheet, headers = _open_student_worksheet(content)
    try:
        mapping = _normalize_mapping(headers, requested_mapping)
        required = {field for field in (required_fields or []) if field in STUDENT_FIELDS}
        missing_mapping = required - set(mapping.values())
        if missing_mapping:
            raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail=f"必填字段未映射: {'、'.join(sorted(missing_mapping))}")
        result: dict[str, Any] = {
            "sheet_name": worksheet.title,
            "columns": [
                {"column": get_column_letter(index), "header": _display_import_value(header) or f"第 {index} 列", "field": mapping.get(index)}
                for index, header in enumerate(headers, start=1)
            ],
            "total_rows": 0,
            "valid_rows": 0,
            "new_rows": 0,
            "existing_rows": 0,
            "conflict_rows": 0,
            "unchanged_rows": 0,
            "skipped_rows": 0,
            "invalid_rows": 0,
            "duplicate_rows": 0,
            "samples": [],
            "conflicts": [],
            "issues": [],
            "required_fields": sorted(required),
            "update_policy": update_policy,
        }
        seen_student_numbers: set[str] = set()
        for row_index, values in enumerate(worksheet.iter_rows(min_row=2, values_only=True), start=2):
            if not any(value is not None and str(value).strip() for value in values):
                continue
            result["total_rows"] += 1
            try:
                raw_fields, parsed_fields = _parsed_excel_fields(values, mapping)
            except ValueError as exc:
                result["invalid_rows"] += 1
                if len(result["issues"]) < 20:
                    result["issues"].append({"row": row_index, "message": str(exc)})
                continue
            student_no = _as_text(raw_fields.get("student_no"))
            full_name = _as_text(raw_fields.get("full_name"))
            if not student_no:
                result["invalid_rows"] += 1
                if len(result["issues"]) < 20:
                    result["issues"].append({"row": row_index, "message": "学号为空"})
                continue
            if student_no in seen_student_numbers:
                result["duplicate_rows"] += 1
                if len(result["issues"]) < 20:
                    result["issues"].append({"row": row_index, "message": "同一文件中学号重复"})
                continue
            seen_student_numbers.add(student_no)
            student = db.scalar(select(Student).where(Student.student_no == student_no))
            if not student and not full_name:
                result["invalid_rows"] += 1
                if len(result["issues"]) < 20:
                    result["issues"].append({"row": row_index, "message": "新学生的姓名不能为空"})
                continue
            row_missing = [field for field in required if not parsed_fields.get(field)]
            if row_missing:
                result["invalid_rows"] += 1
                if len(result["issues"]) < 20:
                    result["issues"].append({"row": row_index, "message": f"必填字段为空: {'、'.join(row_missing)}"})
                continue
            if student and mode == "create_only":
                result["skipped_rows"] += 1
                continue
            if not student and mode == "update_only":
                result["skipped_rows"] += 1
                continue
            result["valid_rows"] += 1
            sample_values = {
                get_column_letter(column): _display_import_value(raw_fields.get(field))
                for column, field in mapping.items()
            }
            if len(result["samples"]) < 8:
                result["samples"].append({"row": row_index, "student_no": student_no, "values": sample_values})
            if student:
                result["existing_rows"] += 1
                changes = []
                for field, incoming in parsed_fields.items():
                    if field == "student_no":
                        continue
                    existing = getattr(student, field)
                    if _display_import_value(existing) != _display_import_value(incoming):
                        changes.append({"field": field, "before": _display_import_value(existing), "after": _display_import_value(incoming)})
                if changes:
                    result["conflict_rows"] += 1
                    if len(result["conflicts"]) < 20:
                        result["conflicts"].append({"row": row_index, "student_no": student_no, "full_name": full_name or student.full_name, "changes": changes})
                else:
                    result["unchanged_rows"] += 1
            else:
                result["new_rows"] += 1
        return result, mapping
    finally:
        workbook.close()


def import_excel(
    db: Session,
    document: SourceDocument,
    content: bytes,
    actor: User,
    mode: str = "upsert",
    requested_mapping: dict[str, Any] | None = None,
    required_fields: list[str] | None = None,
    update_policy: str = "overwrite",
) -> ImportBatch:
    if mode not in IMPORT_MODES:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="无效的导入模式")
    if update_policy not in UPDATE_POLICIES:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="无效的更新策略")
    workbook, worksheet, headers = _open_student_worksheet(content)
    mapping = _normalize_mapping(headers, requested_mapping)
    required = {field for field in (required_fields or []) if field in STUDENT_FIELDS}
    missing_mapping = required - set(mapping.values())
    if missing_mapping:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail=f"必填字段未映射: {'、'.join(sorted(missing_mapping))}")

    batch = ImportBatch(
        source_document_id=document.id,
        imported_by_id=actor.id,
        mode=mode,
        mapping={"columns": {get_column_letter(column): field for column, field in mapping.items()}, "required_fields": sorted(required), "update_policy": update_policy},
    )
    db.add(batch)
    db.flush()
    seen_student_numbers: set[str] = set()
    errors: list[dict[str, Any]] = []
    rollback_data: dict[str, list[dict[str, Any]]] = {"created": [], "updated": []}

    for row_index, values in enumerate(worksheet.iter_rows(min_row=2, values_only=True), start=2):
        if not any(value is not None and str(value).strip() for value in values):
            continue
        batch.total_rows += 1
        try:
            raw_fields, parsed_fields = _parsed_excel_fields(values, mapping)
        except ValueError as exc:
            batch.error_rows += 1
            errors.append({"row": row_index, "message": str(exc)})
            continue
        student_no = _as_text(raw_fields.get("student_no"))
        full_name = _as_text(raw_fields.get("full_name"))
        if not student_no:
            batch.error_rows += 1
            errors.append({"row": row_index, "message": "学号为空"})
            continue
        if student_no in seen_student_numbers:
            batch.error_rows += 1
            errors.append({"row": row_index, "message": "同一文件中学号重复"})
            continue
        seen_student_numbers.add(student_no)

        student = db.scalar(select(Student).where(Student.student_no == student_no))
        if not student and mode == "update_only":
            batch.skipped_rows += 1
            continue
        if student and mode == "create_only":
            batch.skipped_rows += 1
            continue
        if not student and not full_name:
            batch.error_rows += 1
            errors.append({"row": row_index, "message": "新学生的姓名不能为空"})
            continue
        row_missing = [field for field in required if not parsed_fields.get(field)]
        if row_missing:
            batch.error_rows += 1
            errors.append({"row": row_index, "message": f"必填字段为空: {'、'.join(row_missing)}"})
            continue

        if student:
            before_snapshot = student_to_dict(student)
            changed_fields = []
            for field, value in parsed_fields.items():
                if field != "student_no" and (update_policy == "overwrite" or getattr(student, field) in (None, "")):
                    if getattr(student, field) != value:
                        setattr(student, field, value)
                        changed_fields.append(field)
            if changed_fields:
                student.row_version += 1
                db.flush()
                record_student_version(db, student, actor, changed_fields)
                rollback_data["updated"].append(
                    {
                        "student_id": student.id,
                        "student_no": student.student_no,
                        "before": before_snapshot,
                        "expected_row_version": student.row_version,
                        "changed_fields": changed_fields,
                    }
                )
            batch.updated_rows += 1
        else:
            student = Student(**{field: parsed_fields.get(field) for field in STUDENT_FIELDS if field in parsed_fields})
            db.add(student)
            db.flush()
            record_student_version(db, student, actor, list(parsed_fields))
            rollback_data["created"].append(
                {"student_id": student.id, "student_no": student.student_no, "expected_row_version": student.row_version}
            )
            batch.created_rows += 1

        db.flush()
        for column, field in mapping.items():
            raw_value = raw_fields[field]
            if _as_text(raw_value) is None:
                continue
            cell_address = f"{get_column_letter(column)}{row_index}"
            db.add(
                FieldProvenance(
                    student_id=student.id,
                    source_document_id=document.id,
                    import_batch_id=batch.id,
                    field_name=field,
                    source_sheet=worksheet.title,
                    source_row=row_index,
                    source_column=get_column_letter(column),
                    source_locator=cell_address,
                    raw_value=_as_text(raw_value),
                    confidence=100,
                )
            )

    # Keep enough detail for the downloadable correction workbook while the UI
    # still limits the hover report to a compact first page.
    batch.errors = errors[:10000]
    batch.rollback_data = rollback_data if rollback_data["created"] or rollback_data["updated"] else None
    batch.rollback_status = "available" if batch.rollback_data else "not_available"
    batch.status = ImportStatus.COMPLETED_WITH_ERRORS if errors else ImportStatus.COMPLETED
    batch.completed_at = utcnow()
    db.flush()
    workbook.close()
    return batch


def import_word_for_review(db: Session, document: SourceDocument, content: bytes, actor: User) -> list[WordImportCandidate]:
    segments = _word_segments(content)

    candidates = extract_word_candidates(segments)
    stored_candidates: list[WordImportCandidate] = []
    for candidate in candidates:
        data = {field: candidate.get(field) for field in STUDENT_FIELDS if candidate.get(field) is not None}
        if not data.get("student_no") or not data.get("full_name"):
            continue
        candidate_record = WordImportCandidate(
            source_document_id=document.id,
            created_by_id=actor.id,
            candidate_data=data,
            evidence=candidate.get("evidence", []),
            confidence=max(0, min(int(candidate.get("confidence", 0)), 100)),
            status=CandidateStatus.PENDING,
        )
        db.add(candidate_record)
        stored_candidates.append(candidate_record)
    db.flush()
    return stored_candidates


def _word_segments(content: bytes) -> list[dict[str, str]]:
    try:
        word_document = WordDocument(BytesIO(content))
    except Exception as exc:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="无法读取 Word 文件") from exc

    segments: list[dict[str, Any]] = []
    for index, paragraph in enumerate(word_document.paragraphs, start=1):
        if paragraph.text.strip():
            segments.append({"locator": f"段落 {index}", "text": paragraph.text.strip()})
    for table_index, table in enumerate(word_document.tables, start=1):
        for row_index, row in enumerate(table.rows, start=1):
            for column_index, cell in enumerate(row.cells, start=1):
                if cell.text.strip():
                    segments.append({"locator": f"表格 {table_index} 行 {row_index} 列 {column_index}", "text": cell.text.strip()})
    return segments


def _excel_cell_display_value(cell: Any) -> str:
    value = cell.value
    if value is None:
        return ""
    if isinstance(value, (datetime, date)):
        return value.isoformat()
    if isinstance(value, (int, float)) and not isinstance(value, bool):
        number_format = str(cell.number_format or "")
        decimal_match = re.search(r"\.([0#]+)", number_format)
        if decimal_match:
            precision = len(decimal_match.group(1))
            return f"{float(value):.{precision}f}"
        if isinstance(value, float) and value.is_integer():
            return str(int(value))
    return str(value).strip()


def _excel_record_payload(worksheet: Any, header_row_index: int, data_start: int, data_row_index: int) -> dict[str, Any]:
    column_count = worksheet.max_column
    header_rows = [
        [_excel_cell_display_value(cell) for cell in row]
        for row in worksheet.iter_rows(min_row=header_row_index, max_row=data_start - 1, max_col=column_count)
    ]
    data_row = [
        _excel_cell_display_value(cell)
        for cell in next(worksheet.iter_rows(min_row=data_row_index, max_row=data_row_index, max_col=column_count))
    ]
    merged_ranges = []
    for cell_range in worksheet.merged_cells.ranges:
        if cell_range.min_row < header_row_index or cell_range.max_row >= data_start:
            continue
        merged_ranges.append(
            {
                "start_row": cell_range.min_row - header_row_index + 1,
                "end_row": cell_range.max_row - header_row_index + 1,
                "start_column": cell_range.min_col,
                "end_column": cell_range.max_col,
            }
        )
    column_widths = []
    for column_index in range(1, column_count + 1):
        width = worksheet.column_dimensions[get_column_letter(column_index)].width
        column_widths.append(float(width) if width else None)
    return {
        "sheet_name": worksheet.title,
        "header_rows": header_rows,
        "data_row": data_row,
        "merged_ranges": merged_ranges,
        "column_widths": column_widths,
        "source_row": data_row_index,
    }


def _excel_related_segments(content: bytes) -> list[dict[str, Any]]:
    try:
        workbook = load_workbook(BytesIO(content), data_only=True)
    except Exception as exc:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="无法读取 Excel 文件") from exc

    segments: list[dict[str, Any]] = []
    normalized_aliases = {field: {_normalize_header(alias) for alias in aliases} for field, aliases in HEADER_ALIASES.items()}
    worksheet = next((sheet for sheet in workbook.worksheets if sheet.sheet_state == "visible"), None)
    if worksheet is None:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Excel 没有可读取的非隐藏工作表")

    header_row_index = None
    headers = None
    for row_index, row in enumerate(worksheet.iter_rows(min_row=1, max_row=min(12, worksheet.max_row), values_only=True), start=1):
        normalized = {_normalize_header(value) for value in row if _as_text(value)}
        if normalized & normalized_aliases["student_no"] and normalized & normalized_aliases["full_name"]:
            header_row_index = row_index
            headers = row
            break
    if not headers or header_row_index is None:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="第一个非隐藏工作表缺少“学号”和“姓名”表头")

    labels = [_as_text(value) or f"第 {index} 列" for index, value in enumerate(headers, start=1)]
    student_column = next((index for index, label in enumerate(labels) if _normalize_header(label) in normalized_aliases["student_no"]), None)
    name_column = next((index for index, label in enumerate(labels) if _normalize_header(label) in normalized_aliases["full_name"]), None)
    if student_column is None and name_column is None:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="第一个非隐藏工作表未识别到学生标识列")

    data_start = None
    for row_index in range(header_row_index + 1, worksheet.max_row + 1):
        values = [cell.value for cell in next(worksheet.iter_rows(min_row=row_index, max_row=row_index, values_only=False))]
        student_value = _as_text(values[student_column]) if student_column is not None and student_column < len(values) else None
        name_value = _as_text(values[name_column]) if name_column is not None and name_column < len(values) else None
        if student_value or name_value:
            data_start = row_index
            break
        if row_index <= header_row_index + 4:
            for index, value in enumerate(values):
                subheader = _as_text(value)
                if subheader:
                    labels[index] = f"{labels[index]} {subheader}".strip()
    if data_start is None:
        return []

    for row_index, cells in enumerate(worksheet.iter_rows(min_row=data_start, values_only=False), start=data_start):
        values = [cell.value for cell in cells]
        pairs = [f"{labels[index]}：{_excel_cell_display_value(cell)}" for index, cell in enumerate(cells) if _excel_cell_display_value(cell)]
        student_no = _as_text(values[student_column]) if student_column is not None and student_column < len(values) else None
        full_name = _as_text(values[name_column]) if name_column is not None and name_column < len(values) else None
        if not pairs or not (student_no or full_name):
            continue
        identity_columns = {column for column in (student_column, name_column) if column is not None}
        fallback_pairs = [
            f"{labels[index]}：{_excel_cell_display_value(cell)}"
            for index, cell in enumerate(cells)
            if _excel_cell_display_value(cell)
            and index not in identity_columns
            and not any(term in _normalize_header(labels[index]) for term in {"编号", "姓名拼音", "手机", "电话"})
        ]
        segments.append(
            {
                "locator": f"{worksheet.title} 第 {row_index} 行",
                "text": "；".join(pairs),
                "source_sheet": worksheet.title,
                "source_row": str(row_index),
                "student_no": student_no or "",
                "full_name": full_name or "",
                "fallback_remarks": "；".join(fallback_pairs),
                "excel_payload": _excel_record_payload(worksheet, header_row_index, data_start, row_index),
            }
        )
    return segments


def _find_related_student(db: Session, candidate: dict[str, str]) -> tuple[Student | None, str | None]:
    identity_fields = ("student_no", "candidate_no")
    for field in identity_fields:
        value = candidate.get(field)
        if not value:
            continue
        student = db.scalar(select(Student).where(getattr(Student, field) == value))
        if student:
            return student, None
    national_id = candidate.get("national_id")
    if national_id:
        student = db.scalar(select(Student).where(Student.national_id_hash == blind_index(national_id)))
        if student:
            return student, None
    full_name = candidate.get("full_name")
    if not full_name:
        return None, "未识别学生标识"
    matches = list(db.scalars(select(Student).where(Student.full_name == full_name).limit(2)))
    if len(matches) == 1:
        return matches[0], None
    if len(matches) > 1:
        return None, f"姓名“{full_name}”存在重名，缺少唯一标识"
    return None, f"未找到学生“{full_name}”"


def _manual_match_candidates(db: Session, candidate: dict[str, Any]) -> list[int]:
    """Return harmless candidate ids to speed up a reviewer, never auto-apply a match."""
    student_no = str(candidate.get("student_no") or "").strip()
    candidate_no = str(candidate.get("candidate_no") or "").strip()
    full_name = str(candidate.get("full_name") or "").strip()
    statement = select(Student.id)
    if student_no:
        return list(db.scalars(statement.where(Student.student_no == student_no).limit(8)))
    if candidate_no:
        return list(db.scalars(statement.where(Student.candidate_no == candidate_no).limit(8)))
    if full_name:
        return list(db.scalars(statement.where(Student.full_name == full_name).limit(8)))
    return []


def _make_related_info_candidate(
    *,
    document: SourceDocument,
    batch: ImportBatch,
    student: Student,
    candidate: dict[str, Any],
    source: dict[str, Any],
) -> RelatedInfoCandidate:
    return RelatedInfoCandidate(
        source_document_id=document.id,
        import_batch_id=batch.id,
        student_id=student.id,
        extracted_identity={field: value for field, value in candidate.items() if field in {"student_no", "candidate_no", "national_id", "full_name"} and value},
        remarks=str(candidate.get("remarks") or ""),
        content_type=str(candidate.get("content_type") or "text"),
        excel_payload=candidate.get("excel_payload"),
        source_sheet=source.get("source_sheet"),
        source_row=int(source["source_row"]) if source.get("source_row") else None,
        source_locator=str(candidate.get("locator") or source.get("locator") or "导入记录")[:255],
        confidence=max(0, min(int(candidate.get("confidence") or 85), 100)),
        status=CandidateStatus.PENDING,
    )


def _append_remarks(existing: str | None, incoming: str) -> tuple[str, bool]:
    existing = (existing or "").strip()
    incoming = incoming.strip()[:1200]
    if not incoming or incoming in existing:
        return existing, False
    if not existing:
        return incoming, True
    available = 2000 - len(incoming) - 2
    if available <= 0:
        return incoming[:2000], True
    return f"{existing[-available:].strip()}\n\n{incoming}", True


def _merge_candidate_source(candidate: RelatedInfoCandidate, source: dict[str, str], locator: str) -> None:
    locations = [item.strip() for item in candidate.source_locator.split("；") if item.strip()]
    if locator not in locations:
        locations.append(locator)
        candidate.source_locator = "；".join(locations)[:255]
    source_sheet = source.get("source_sheet")
    if source_sheet and candidate.source_sheet and candidate.source_sheet != source_sheet:
        candidate.source_sheet = "多工作表"
        candidate.source_row = None
    elif source_sheet and not candidate.source_sheet:
        candidate.source_sheet = source_sheet
    elif candidate.source_row != (int(source["source_row"]) if source.get("source_row") else None):
        candidate.source_row = None


def _excel_related_candidates(segments: list[dict[str, Any]]) -> list[dict[str, Any]]:
    return [
        {
            "student_no": segment.get("student_no", ""),
            "candidate_no": "",
            "national_id": "",
            "full_name": segment.get("full_name", ""),
            "remarks": "Excel 原始行记录",
            "locator": segment["locator"],
            "confidence": 100,
            "content_type": "excel_card",
            "excel_payload": segment["excel_payload"],
        }
        for segment in segments
    ]


def import_related_info(db: Session, document: SourceDocument, content: bytes, actor: User) -> ImportBatch:
    if document.file_type == "word":
        segments = _word_segments(content)
        candidates = extract_related_info(segments)
        analysis_type = "local_ai"
    elif document.file_type == "excel":
        segments = _excel_related_segments(content)
        candidates = _excel_related_candidates(segments)
        analysis_type = "structured_excel"
    else:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="学生相关信息仅支持 Word 或 Excel 文件")
    if not segments:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="文件中没有可分析的文字或表格内容")

    batch = ImportBatch(
        source_document_id=document.id,
        imported_by_id=actor.id,
        mode="related_info",
        mapping={"analysis": analysis_type, "file_type": document.file_type},
        total_rows=len(segments),
    )
    db.add(batch)
    db.flush()
    segment_index = {segment["locator"]: segment for segment in segments}
    errors: list[dict[str, Any]] = []
    pending_by_student: dict[int, RelatedInfoCandidate] = {}

    if not candidates:
        batch.error_rows = 1
        errors.append({"row": "AI", "message": "本地 AI 未识别出可写入备注的学生相关信息"})
    for candidate in candidates:
        student, error = _find_related_student(db, candidate)
        if not student:
            batch.skipped_rows += 1
            errors.append({"row": candidate["locator"], "message": error})
            source = segment_index.get(candidate["locator"], {})
            db.add(
                ImportMatchReview(
                    source_document_id=document.id,
                    import_batch_id=batch.id,
                    identity={field: value for field, value in candidate.items() if field in {"student_no", "candidate_no", "national_id", "full_name"} and value},
                    payload={
                        "remarks": candidate.get("remarks"),
                        "content_type": candidate.get("content_type", "text"),
                        "excel_payload": candidate.get("excel_payload"),
                        "locator": candidate.get("locator"),
                        "confidence": candidate.get("confidence", 85),
                        "source_sheet": source.get("source_sheet"),
                        "source_row": source.get("source_row"),
                    },
                    candidate_student_ids=_manual_match_candidates(db, candidate),
                    match_reason=error,
                )
            )
            continue
        existing = pending_by_student.get(student.id)
        if existing:
            if candidate.get("content_type") == "excel_card":
                batch.skipped_rows += 1
                continue
            remarks, changed = _append_remarks(existing.remarks, candidate["remarks"])
            if not changed:
                batch.skipped_rows += 1
                continue
            existing.remarks = remarks
            _merge_candidate_source(existing, segment_index.get(candidate["locator"], {}), candidate["locator"])
            continue
        source = segment_index.get(candidate["locator"], {})
        pending = _make_related_info_candidate(document=document, batch=batch, student=student, candidate=candidate, source=source)
        db.add(pending)
        pending_by_student[student.id] = pending
        batch.created_rows += 1

    batch.errors = errors[:200]
    batch.error_rows = len(errors)
    batch.status = ImportStatus.COMPLETED_WITH_ERRORS if errors else ImportStatus.COMPLETED
    batch.completed_at = utcnow()
    db.flush()
    batch.rollback_data = {"related_changes": []}
    batch.rollback_status = "available"
    db.flush()
    return batch


def resolve_import_match_review(db: Session, review: ImportMatchReview, student: Student, actor: User) -> RelatedInfoCandidate:
    if review.status != "pending":
        raise HTTPException(status_code=status.HTTP_409_CONFLICT, detail="该人工匹配项已处理")
    document = db.get(SourceDocument, review.source_document_id)
    batch = db.get(ImportBatch, review.import_batch_id)
    if not document or not batch:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="关联的导入来源不存在")
    if db.scalar(select(RelatedInfoCandidate).where(RelatedInfoCandidate.import_batch_id == batch.id, RelatedInfoCandidate.student_id == student.id)):
        raise HTTPException(status_code=status.HTTP_409_CONFLICT, detail="该学生已有同一批次的待审核信息")
    payload = review.payload or {}
    candidate = dict(review.identity or {}) | {
        "remarks": payload.get("remarks") or "",
        "content_type": payload.get("content_type") or "text",
        "excel_payload": payload.get("excel_payload"),
        "locator": payload.get("locator") or "人工匹配记录",
        "confidence": payload.get("confidence") or 80,
    }
    pending = _make_related_info_candidate(document=document, batch=batch, student=student, candidate=candidate, source=payload)
    db.add(pending)
    review.status = "matched"
    review.matched_student_id = student.id
    review.reviewed_by_id = actor.id
    review.reviewed_at = utcnow()
    return pending


def dismiss_import_match_review(review: ImportMatchReview, actor: User) -> None:
    if review.status != "pending":
        raise HTTPException(status_code=status.HTTP_409_CONFLICT, detail="该人工匹配项已处理")
    review.status = "ignored"
    review.reviewed_by_id = actor.id
    review.reviewed_at = utcnow()


def apply_related_info_candidate(db: Session, candidate: RelatedInfoCandidate, actor: User) -> Student:
    if candidate.status != CandidateStatus.PENDING:
        raise HTTPException(status_code=status.HTTP_409_CONFLICT, detail="该审核项已处理")
    student = db.get(Student, candidate.student_id)
    if not student:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="关联学生不存在")
    batch = db.get(ImportBatch, candidate.import_batch_id)
    if not batch:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="导入批次不存在")
    rollback_data = dict(batch.rollback_data or {})
    related_changes = list(rollback_data.get("related_changes") or [])
    if candidate.content_type == "excel_card":
        if not candidate.excel_payload:
            raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Excel 原始记录缺失")
        document = db.get(SourceDocument, candidate.source_document_id)
        if not document:
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Excel 导入来源不存在")
        card = StudentRelatedInfoCard(
            student_id=student.id,
            source_document_id=document.id,
            import_batch_id=batch.id,
            imported_by_id=batch.imported_by_id,
            title=Path(document.original_filename).stem,
            excel_payload=candidate.excel_payload,
            imported_at=batch.created_at,
        )
        db.add(card)
        student.row_version += 1
        db.flush()
        record_student_version(db, student, actor, ["remarks"])
        db.add(
            FieldProvenance(
                student_id=student.id,
                source_document_id=candidate.source_document_id,
                import_batch_id=candidate.import_batch_id,
                field_name="remarks",
                source_sheet=candidate.source_sheet,
                source_row=candidate.source_row,
                source_locator=candidate.source_locator,
                raw_value=Path(document.original_filename).stem,
                confidence=candidate.confidence,
            )
        )
        related_changes.append({"kind": "excel_card", "candidate_id": candidate.id, "card_id": card.id, "student_id": student.id, "student_no": student.student_no})
    else:
        before_remarks = student.remarks or ""
        remarks, changed = _append_remarks(student.remarks, candidate.remarks)
        if changed:
            student.remarks = remarks
            student.row_version += 1
            record_student_version(db, student, actor, ["remarks"])
            db.add(
                FieldProvenance(
                    student_id=student.id,
                    source_document_id=candidate.source_document_id,
                    import_batch_id=candidate.import_batch_id,
                    field_name="remarks",
                    source_sheet=candidate.source_sheet,
                    source_row=candidate.source_row,
                    source_locator=candidate.source_locator,
                    raw_value=candidate.remarks,
                    confidence=candidate.confidence,
                )
            )
            related_changes.append({"kind": "remarks", "candidate_id": candidate.id, "student_id": student.id, "student_no": student.student_no, "before_remarks": before_remarks, "after_remarks": remarks})
    candidate.status = CandidateStatus.APPROVED
    candidate.reviewed_by_id = actor.id
    candidate.reviewed_at = utcnow()
    batch.rollback_data = rollback_data | {"related_changes": related_changes}
    db.flush()
    return student


def apply_word_candidate(db: Session, candidate: WordImportCandidate, payload: CandidateApproval, actor: User | None = None) -> Student:
    if candidate.status != CandidateStatus.PENDING:
        raise HTTPException(status_code=status.HTTP_409_CONFLICT, detail="该候选项已处理")
    values = payload.model_dump()
    student = db.scalar(select(Student).where(Student.student_no == payload.student_no))
    if student:
        changed_fields = []
        for field, value in values.items():
            if value is not None and field != "student_no":
                if getattr(student, field) != value:
                    setattr(student, field, value)
                    changed_fields.append(field)
        if changed_fields:
            student.row_version += 1
            record_student_version(db, student, actor, changed_fields)
    else:
        student = Student(**values)
        db.add(student)
        db.flush()
        record_student_version(db, student, actor, list(values))
    for evidence in candidate.evidence:
        field = evidence.get("field")
        if field not in STUDENT_FIELDS:
            continue
        db.add(
            FieldProvenance(
                student_id=student.id,
                source_document_id=candidate.source_document_id,
                field_name=field,
                source_locator=evidence.get("locator", "Word 文档"),
                raw_value=str(evidence.get("value", "")),
                confidence=candidate.confidence,
            )
        )
    candidate.status = CandidateStatus.APPROVED
    db.flush()
    return student

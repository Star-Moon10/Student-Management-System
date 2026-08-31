"""Create a student-scoped copy of a source document without altering its original binary."""

from collections.abc import Iterable
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document as WordDocument
from openpyxl import load_workbook
from openpyxl.cell.cell import MergedCell

from app.models import Student


REDACTED_VALUE = "【其他学生信息已脱敏】"


def sanitized_download_filename(original_filename: str, student_no: str) -> str:
    source = Path(original_filename)
    return f"{source.stem}_仅{student_no}学生可见{source.suffix}"


def _student_tokens(student: Student) -> set[str]:
    return {
        str(value).strip()
        for value in (
            student.student_no,
            student.candidate_no,
            student.full_name,
            student.national_id,
            student.mobile_phone,
            student.electronic_email,
        )
        if value is not None and str(value).strip()
    }


def _contains_any(value: Any, tokens: Iterable[str]) -> bool:
    text = str(value or "")
    return any(token in text for token in tokens)


def _replace_tokens(value: Any, tokens: Iterable[str]) -> Any:
    if not isinstance(value, str):
        return value
    result = value
    for token in sorted(set(tokens), key=len, reverse=True):
        result = result.replace(token, "***")
    return result


def _sanitize_excel(
    content: bytes,
    target: Student,
    students: Iterable[Student],
    known_rows: dict[str, dict[int, int]],
) -> bytes:
    workbook = load_workbook(BytesIO(content))
    target_tokens = _student_tokens(target)
    other_tokens = set().union(*(_student_tokens(student) for student in students if student.id != target.id))
    for worksheet in workbook.worksheets:
        sheet_rows = known_rows.get(worksheet.title, {})
        for row_index, row in enumerate(worksheet.iter_rows(), start=1):
            values = [cell.value for cell in row if not isinstance(cell, MergedCell)]
            known_student_id = sheet_rows.get(row_index)
            row_is_target = known_student_id == target.id or _contains_any(" ".join(str(value or "") for value in values), target_tokens)
            row_is_other = (known_student_id is not None and known_student_id != target.id) or _contains_any(" ".join(str(value or "") for value in values), other_tokens)
            if row_is_other and not row_is_target:
                for cell in row:
                    if not isinstance(cell, MergedCell) and cell.value not in (None, ""):
                        cell.value = REDACTED_VALUE
                continue
            if row_is_target:
                for cell in row:
                    if not isinstance(cell, MergedCell):
                        cell.value = _replace_tokens(cell.value, other_tokens)
    stream = BytesIO()
    workbook.save(stream)
    return stream.getvalue()


def _sanitize_word(
    content: bytes,
    target: Student,
    students: Iterable[Student],
) -> bytes:
    document = WordDocument(BytesIO(content))
    target_tokens = _student_tokens(target)
    other_tokens = set().union(*(_student_tokens(student) for student in students if student.id != target.id))

    def sanitize_text(value: str) -> str:
        has_target = _contains_any(value, target_tokens)
        has_other = _contains_any(value, other_tokens)
        if has_other and not has_target:
            return REDACTED_VALUE
        return _replace_tokens(value, other_tokens)

    seen_cells: set[int] = set()
    for paragraph in document.paragraphs:
        paragraph.text = sanitize_text(paragraph.text)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if id(cell._tc) in seen_cells:
                    continue
                seen_cells.add(id(cell._tc))
                for paragraph in cell.paragraphs:
                    paragraph.text = sanitize_text(paragraph.text)
    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()


def create_student_scoped_source_copy(
    content: bytes,
    file_type: str,
    target: Student,
    students: Iterable[Student],
    known_rows: dict[str, dict[int, int]] | None = None,
) -> bytes:
    if file_type == "excel":
        return _sanitize_excel(content, target, students, known_rows or {})
    if file_type == "word":
        return _sanitize_word(content, target, students)
    raise ValueError("不支持生成该类型的学生范围副本")
